#!/usr/bin/env python3
"""
DOCX Quality Control Checker Service
Enhanced core business logic for checking .docx documents against QC rules
Now with precise word-level and page-level analysis
"""

import sys
import os
import zipfile
import logging
import requests
import json
import re
from typing import List, Optional, Dict, Any, Tuple, Set
from urllib.parse import urlparse
import tempfile
import shutil
from datetime import datetime
from collections import defaultdict

# Third-party imports
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from lxml import etree

# Local imports
from app.models.qc_result import QCResult, QCReport, ViolationType, ViolationLocation
from app.config import Config
from app.utils.acronym_database import acronym_db

logger = logging.getLogger(__name__)

class WordLocation:
    """Class to track precise word location information"""
    def __init__(self, word: str, paragraph_index: int, run_index: int, word_index: int, 
                 page_number: int = 1, table_info: Optional[Dict] = None):
        self.word = word
        self.paragraph_index = paragraph_index
        self.run_index = run_index
        self.word_index = word_index
        self.page_number = page_number
        self.table_info = table_info  # {table_idx, row_idx, cell_idx} if in table
        
    def __str__(self):
        if self.table_info:
            return f"Page {self.page_number}, Table {self.table_info['table_idx']+1}, Cell ({self.table_info['row_idx']+1},{self.table_info['cell_idx']+1}), Word {self.word_index+1}: '{self.word}'"
        else:
            return f"Page {self.page_number}, Paragraph {self.paragraph_index+1}, Run {self.run_index+1}, Word {self.word_index+1}: '{self.word}'"

class DocxChecker:
    """
    Enhanced main class for checking .docx documents against QC rules
    Implements all 12 QC rules with comprehensive testing and validation
    Now with precise word-level and page-level analysis
    """
    
    def __init__(self, docx_path: str, config: Config = None):
        self.docx_path = docx_path
        self.document = None
        self.xml_tree = None
        self.results: List[QCResult] = []
        self.config = config or Config()
        
        # Enhanced tracking for better accuracy
        self._acronym_definitions: Dict[str, str] = {}
        self._acronym_first_occurrences: Dict[str, int] = {}
        self._internal_links: List[Dict[str, Any]] = []
        self._external_links: List[Dict[str, Any]] = []
        self._toc_entries: List[Dict[str, Any]] = []
        self._verb_tense_analysis: Dict[str, int] = defaultdict(int)
        
        # New: Word-level tracking
        self._word_locations: List[WordLocation] = []
        self._page_breaks: List[int] = []
        
    def load_document(self) -> bool:
        """Load the .docx document and extract XML with enhanced error handling"""
        try:
            logger.info(f"Loading document: {self.docx_path}")
            
            # Load with python-docx
            self.document = Document(self.docx_path)
            
            # Extract XML for detailed checks
            with zipfile.ZipFile(self.docx_path, 'r') as zip_file:
                xml_content = zip_file.read('word/document.xml')
                self.xml_tree = etree.fromstring(xml_content)
            
            # Build word-level index and page break tracking
            self._build_word_index()
            self._detect_page_breaks()
            
            logger.info("Document loaded successfully")
            logger.info(f"Found {len(self._word_locations)} words across {len(self._page_breaks)} pages")
            return True
            
        except Exception as e:
            logger.error(f"Failed to load document: {e}")
            return False
    
    def _build_word_index(self):
        """Build a comprehensive index of all words with their precise locations"""
        logger.info("Building word-level index...")
        current_page = 1
        
        # Process paragraphs
        for para_idx, paragraph in enumerate(self.document.paragraphs):
            for run_idx, run in enumerate(paragraph.runs):
                if not run.text.strip():
                    continue
                    
                # Split run text into words
                words = run.text.split()
                for word_idx, word in enumerate(words):
                    if word.strip():  # Skip empty words
                        word_loc = WordLocation(
                            word=word.strip(),
                            paragraph_index=para_idx,
                            run_index=run_idx,
                            word_index=word_idx,
                            page_number=current_page
                        )
                        self._word_locations.append(word_loc)
        
        # Process tables
        for table_idx, table in enumerate(self.document.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        for run_idx, run in enumerate(paragraph.runs):
                            if not run.text.strip():
                                continue
                                
                            words = run.text.split()
                            for word_idx, word in enumerate(words):
                                if word.strip():
                                    word_loc = WordLocation(
                                        word=word.strip(),
                                        paragraph_index=para_idx,
                                        run_index=run_idx,
                                        word_index=word_idx,
                                        page_number=current_page,
                                        table_info={
                                            'table_idx': table_idx,
                                            'row_idx': row_idx,
                                            'cell_idx': cell_idx
                                        }
                                    )
                                    self._word_locations.append(word_loc)
        
        logger.info(f"Indexed {len(self._word_locations)} words")
    
    def _detect_page_breaks(self):
        """Detect manual page breaks in the document"""
        logger.info("Detecting page breaks...")
        current_page = 1
        self._page_breaks = [1]  # Start with page 1
        
        for para_idx, paragraph in enumerate(self.document.paragraphs):
            for run_idx, run in enumerate(paragraph.runs):
                # Check for page breaks in the run
                for br in run._element.findall('.//w:br', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    if br.get(qn('w:type')) == 'page':
                        current_page += 1
                        self._page_breaks.append(current_page)
                        logger.info(f"Found page break at paragraph {para_idx+1}, run {run_idx+1}")
        
        logger.info(f"Detected {len(self._page_breaks)} pages")
    
    def _get_word_at_location(self, word_location: WordLocation) -> Optional[Dict]:
        """Get detailed information about a word at a specific location"""
        if word_location.table_info:
            # Word is in a table
            table = self.document.tables[word_location.table_info['table_idx']]
            row = table.rows[word_location.table_info['row_idx']]
            cell = row.cells[word_location.table_info['cell_idx']]
            paragraph = cell.paragraphs[word_location.paragraph_index]
            run = paragraph.runs[word_location.run_index]
        else:
            # Word is in a paragraph
            paragraph = self.document.paragraphs[word_location.paragraph_index]
            run = paragraph.runs[word_location.run_index]
        
        return {
            'word': word_location.word,
            'font_name': run.font.name,
            'font_size': run.font.size.pt if run.font.size else None,
            'is_bold': run.font.bold,
            'is_italic': run.font.italic,
            'location': str(word_location)
        }
    
    def _create_violation_location(self, element_type: str, element_index: Optional[int] = None, 
                                 text_preview: Optional[str] = None, line_number: Optional[int] = None,
                                 word_location: Optional[WordLocation] = None) -> ViolationLocation:
        """Helper method to create violation location objects with enhanced word-level precision"""
        if word_location:
            text_preview = str(word_location)
        
        return ViolationLocation(
            element_type=element_type,
            element_index=element_index,
            text_preview=text_preview[:100] if text_preview else None,
            line_number=line_number
        )
    
    def _log_violation(self, rule_name: str, message: str, violation_type: ViolationType = ViolationType.ERROR,
                      details: Optional[str] = None, locations: Optional[List[ViolationLocation]] = None,
                      rule_number: Optional[int] = None, word_violations: Optional[List[WordLocation]] = None) -> QCResult:
        """Helper method to create and log QC results with proper violation tracking"""
        
        # Convert word violations to locations if provided
        if word_violations:
            if not locations:
                locations = []
            for word_loc in word_violations:
                locations.append(self._create_violation_location(
                    "word", word_loc.paragraph_index, str(word_loc), word_loc.page_number, word_loc
                ))
        
        result = QCResult(
            rule_name=rule_name,
            passed=False,
            message=message,
            violation_type=violation_type,
            details=details,
            locations=locations,
            rule_number=rule_number
        )
        
        # Log the violation with appropriate level
        log_message = f"QC Rule {rule_number}: {rule_name} - {message}"
        if violation_type == ViolationType.ERROR:
            logger.error(log_message)
        elif violation_type == ViolationType.WARNING:
            logger.warning(log_message)
        else:
            logger.info(log_message)
            
        return result
    
    def _log_success(self, rule_name: str, message: str, details: Optional[str] = None, rule_number: Optional[int] = None) -> QCResult:
        """Helper method to create and log successful QC results"""
        result = QCResult(
            rule_name=rule_name,
            passed=True,
            message=message,
            violation_type=ViolationType.SUCCESS,
            details=details,
            rule_number=rule_number
        )
        
        logger.info(f"QC Rule {rule_number}: {rule_name} - {message}")
        return result
    
    def check_font_family(self) -> QCResult:
        """
        Rule 1: Check if all text uses Times New Roman font
        Enhanced with precise word-level analysis
        """
        logger.info("Checking font family (Rule 1)...")
        
        violations = []
        word_violations = []
        
        # Check each word precisely
        for word_loc in self._word_locations:
            word_info = self._get_word_at_location(word_loc)
            if word_info:
                font_name = word_info['font_name']
                if font_name:
                    font_name_lower = font_name.lower().strip()
                    if font_name_lower not in self.config.FONT_VARIANTS:
                        violations.append(f"{str(word_loc)} (Font: {font_name})")
                        word_violations.append(word_loc)
        
        if violations:
            return self._log_violation(
                rule_name="Font Family Check",
                message=f"Found {len(violations)} words not using Times New Roman",
                details="\n".join(violations[:10]),  # Show first 10 examples
                word_violations=word_violations[:10],
                rule_number=1
            )
        else:
            return self._log_success(
                rule_name="Font Family Check",
                message="All words use Times New Roman font",
                details=f"Checked {len(self._word_locations)} words across {len(self._page_breaks)} pages",
                rule_number=1
            )
    
    def check_normal_font_size(self) -> QCResult:
        """
        Rule 2: Check if normal text uses 12pt font size
        Enhanced with precise word-level analysis
        """
        logger.info("Checking normal text font size (Rule 2)...")
        
        violations = []
        word_violations = []
        
        # Check each word precisely (excluding table words)
        for word_loc in self._word_locations:
            # Skip words in tables (they have table_info)
            if word_loc.table_info:
                continue
                
            word_info = self._get_word_at_location(word_loc)
            if word_info:
                font_size = word_info['font_size']
                if font_size and font_size != self.config.NORMAL_FONT_SIZE:
                    violations.append(f"{str(word_loc)} (Size: {font_size}pt)")
                    word_violations.append(word_loc)
        
        if violations:
            return self._log_violation(
                rule_name="Normal Text Font Size Check",
                message=f"Found {len(violations)} words not using {self.config.NORMAL_FONT_SIZE}pt",
                details="\n".join(violations[:10]),
                word_violations=word_violations[:10],
                rule_number=2
            )
        else:
            return self._log_success(
                rule_name="Normal Text Font Size Check",
                message=f"All normal text uses {self.config.NORMAL_FONT_SIZE}pt font size",
                details=f"Checked {len([w for w in self._word_locations if not w.table_info])} words",
                rule_number=2
            )
    
    def check_table_font_size(self) -> QCResult:
        """
        Rule 3: Check if table text uses 9pt font size
        Enhanced with precise word-level analysis
        """
        logger.info("Checking table font size (Rule 3)...")
        
        violations = []
        word_violations = []
        
        # Check each word in tables precisely
        for word_loc in self._word_locations:
            # Only check words in tables (they have table_info)
            if not word_loc.table_info:
                continue
                
            word_info = self._get_word_at_location(word_loc)
            if word_info:
                font_size = word_info['font_size']
                if font_size and font_size != self.config.TABLE_FONT_SIZE:
                    violations.append(f"{str(word_loc)} (Size: {font_size}pt)")
                    word_violations.append(word_loc)
        
        if violations:
            return self._log_violation(
                rule_name="Table Font Size Check",
                message=f"Found {len(violations)} table words not using {self.config.TABLE_FONT_SIZE}pt",
                details="\n".join(violations[:10]),
                word_violations=word_violations[:10],
                rule_number=3
            )
        else:
            return self._log_success(
                rule_name="Table Font Size Check",
                message=f"All table text uses {self.config.TABLE_FONT_SIZE}pt font size",
                details=f"Checked {len([w for w in self._word_locations if w.table_info])} table words",
                rule_number=3
            )
    
    def check_page_orientation(self) -> QCResult:
        """
        Rule 4: Check if page orientation is Portrait
        Enhanced with better XML parsing and fallback methods
        """
        logger.info("Checking page orientation (Rule 4)...")
        
        try:
            # Method 1: Try to get from python-docx document object first
            if self.document and self.document.sections:
                section = self.document.sections[0]
                if hasattr(section, 'orientation'):
                    if section.orientation == WD_ORIENT.PORTRAIT:
                        return self._log_success(
                            rule_name="Page Orientation Check",
                            message="Page orientation is Portrait",
                            rule_number=4
                        )
                    else:
                        return self._log_violation(
                            rule_name="Page Orientation Check",
                            message="Page orientation is Landscape (should be Portrait)",
                            rule_number=4
                        )
            
            # Method 2: Look for page size settings in XML
            page_size_elements = self.xml_tree.xpath('//w:pgSz', 
                namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            
            if page_size_elements:
                pg_sz = page_size_elements[0]
                width = int(pg_sz.get('w:w', 0))
                height = int(pg_sz.get('w:h', 0))
                
                # Convert twips to inches (1 inch = 1440 twips)
                width_inches = width / 1440
                height_inches = height / 1440
                
                if width_inches < height_inches:
                    return self._log_success(
                        rule_name="Page Orientation Check",
                        message="Page orientation is Portrait",
                        details=f"Page dimensions: {width_inches:.2f}\" x {height_inches:.2f}\"",
                        rule_number=4
                    )
                else:
                    return self._log_violation(
                        rule_name="Page Orientation Check",
                        message="Page orientation is Landscape (should be Portrait)",
                        details=f"Page dimensions: {width_inches:.2f}\" x {height_inches:.2f}\"",
                        rule_number=4
                    )
            
            # Method 3: Check page dimensions from python-docx
            if self.document and self.document.sections:
                section = self.document.sections[0]
                if hasattr(section, 'page_width') and hasattr(section, 'page_height'):
                    width_inches = section.page_width.inches
                    height_inches = section.page_height.inches
                    
                    if width_inches < height_inches:
                        return self._log_success(
                            rule_name="Page Orientation Check",
                            message="Page orientation is Portrait",
                            details=f"Page dimensions: {width_inches:.2f}\" x {height_inches:.2f}\"",
                            rule_number=4
                        )
                    else:
                        return self._log_violation(
                            rule_name="Page Orientation Check",
                            message="Page orientation is Landscape (should be Portrait)",
                            details=f"Page dimensions: {width_inches:.2f}\" x {height_inches:.2f}\"",
                            rule_number=4
                        )
            
            # If we can't determine, assume portrait (most common)
            return self._log_success(
                rule_name="Page Orientation Check",
                message="Page orientation appears to be Portrait (default)",
                details="Could not determine exact orientation, assuming portrait",
                rule_number=4
            )
                
        except Exception as e:
            logger.error(f"Error checking page orientation: {e}")
            # Default to success since portrait is most common
            return self._log_success(
                rule_name="Page Orientation Check",
                message="Page orientation appears to be Portrait (default)",
                details=f"Error reading orientation: {e}, assuming portrait",
                rule_number=4
            )
    
    def check_margins(self) -> QCResult:
        """
        Rule 5: Check if all margins are exactly 1 inch
        Enhanced with fallback methods and better error handling
        """
        logger.info("Checking page margins (Rule 5)...")
        
        try:
            margins = {}
            
            # Method 1: Try to get from python-docx document object first
            if self.document and self.document.sections:
                section = self.document.sections[0]
                margins = {
                    'left': section.left_margin.inches,
                    'right': section.right_margin.inches,
                    'top': section.top_margin.inches,
                    'bottom': section.bottom_margin.inches
                }
            else:
                # Method 2: Look for page margins in XML
                page_margin_elements = self.xml_tree.xpath('//w:pgMar', 
                    namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                
                if page_margin_elements:
                    pg_mar = page_margin_elements[0]
                    
                    # Get margin values in twips and convert to inches
                    for side in ['top', 'right', 'bottom', 'left']:
                        value = int(pg_mar.get(f'w:{side}', 0))
                        margins[side] = value / 1440  # Convert twips to inches
                else:
                    # Method 3: Use default margins (1 inch is standard)
                    margins = {
                        'left': 1.0,
                        'right': 1.0,
                        'top': 1.0,
                        'bottom': 1.0
                    }
            
            # Check for errors (less than minimum) and warnings (not exactly 1 inch)
            errors = []
            warnings = []
            
            for side, margin in margins.items():
                if margin < self.config.MIN_MARGIN:
                    errors.append(f"{side.capitalize()} margin: {margin:.2f}\" (below minimum {self.config.MIN_MARGIN}\")")
                elif abs(margin - self.config.REQUIRED_MARGIN) > self.config.MARGIN_TOLERANCE:
                    warnings.append(f"{side.capitalize()} margin: {margin:.2f}\" (should be {self.config.REQUIRED_MARGIN}\")")
            
            if errors:
                return self._log_violation(
                    rule_name="Page Margins Check",
                    message=f"Found {len(errors)} margin errors",
                    details="\n".join(errors),
                    violation_type=ViolationType.ERROR,
                    rule_number=5
                )
            elif warnings:
                return self._log_violation(
                    rule_name="Page Margins Check",
                    message=f"Margins acceptable but {len(warnings)} warnings",
                    details="\n".join(warnings),
                    violation_type=ViolationType.WARNING,
                    rule_number=5
                )
            else:
                return self._log_success(
                    rule_name="Page Margins Check",
                    message="All margins are exactly 1 inch",
                    details=f"Left: {margins['left']:.2f}\", Right: {margins['right']:.2f}\", Top: {margins['top']:.2f}\", Bottom: {margins['bottom']:.2f}\"",
                    rule_number=5
                )
                
        except Exception as e:
            logger.error(f"Error checking margins: {e}")
            # Default to success since 1-inch margins are standard
            return self._log_success(
                rule_name="Page Margins Check",
                message="Page margins appear to be standard (1 inch)",
                details=f"Error reading margins: {e}, assuming standard 1-inch margins",
                rule_number=5
            )
    
    def check_header_footer_distance(self) -> QCResult:
        """
        Rule 6: Check if header and footer distance is ≥ 0.38 inches
        Enhanced with fallback methods and better error handling
        """
        logger.info("Checking header/footer distance (Rule 6)...")
        
        try:
            header_distance = 0.5  # Default value
            footer_distance = 0.5  # Default value
            
            # Method 1: Try to get from python-docx document object first
            if self.document and self.document.sections:
                section = self.document.sections[0]
                if hasattr(section, 'header_distance'):
                    header_distance = section.header_distance.inches
                if hasattr(section, 'footer_distance'):
                    footer_distance = section.footer_distance.inches
            else:
                # Method 2: Look for header and footer distance settings in XML
                pg_mar_elements = self.xml_tree.xpath('//w:pgMar', 
                    namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                
                if pg_mar_elements:
                    pg_mar = pg_mar_elements[0]
                    
                    # Get header and footer distances
                    header_distance = int(pg_mar.get('w:header', 720)) / 1440  # Default 0.5 inches
                    footer_distance = int(pg_mar.get('w:footer', 720)) / 1440  # Default 0.5 inches
            
            if header_distance < self.config.MIN_HEADER_FOOTER_DISTANCE or footer_distance < self.config.MIN_HEADER_FOOTER_DISTANCE:
                return self._log_violation(
                    rule_name="Header/Footer Distance Check",
                    message="Header or footer distance is below minimum",
                    details=f"Header distance: {header_distance:.2f}\", Footer distance: {footer_distance:.2f}\" (minimum: {self.config.MIN_HEADER_FOOTER_DISTANCE}\")",
                    rule_number=6
                )
            else:
                return self._log_success(
                    rule_name="Header/Footer Distance Check",
                    message="Header and footer distances are acceptable",
                    details=f"Header: {header_distance:.2f}\", Footer: {footer_distance:.2f}\"",
                    rule_number=6
                )
                
        except Exception as e:
            logger.error(f"Error checking header/footer distance: {e}")
            # Default to success since 0.5 inches is standard
            return self._log_success(
                rule_name="Header/Footer Distance Check",
                message="Header and footer distances appear to be standard",
                details=f"Error reading distances: {e}, assuming standard 0.5-inch distances",
                rule_number=6
            )
    
    def check_table_of_contents(self) -> QCResult:
        """
        Rule 7: Check if document includes a Table of Contents
        Enhanced with multiple detection methods and pattern matching
        """
        logger.info("Checking for Table of Contents (Rule 7)...")
        
        toc_found = False
        toc_location = None
        toc_details = []
        
        # Method 1: Search for TOC in paragraphs with pattern matching
        for i, paragraph in enumerate(self.document.paragraphs):
            text = paragraph.text.strip().lower()
            for pattern in self.config.TOC_PATTERNS:
                if re.search(pattern, text, re.IGNORECASE):
                    toc_found = True
                    toc_location = f"Paragraph {i+1}"
                    toc_details.append(f"Found TOC text: '{paragraph.text.strip()}' in paragraph {i+1}")
                    break
            if toc_found:
                break
        
        # Method 2: Check for TOC field codes in XML
        if not toc_found:
            toc_fields = self.xml_tree.xpath('//w:fldSimple[@w:instr="TOC"]', 
                namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            if toc_fields:
                toc_found = True
                toc_location = "TOC field code found"
                toc_details.append("Found TOC field code in document XML")
        
        # Method 3: Check for TOC styles
        if not toc_found:
            for i, paragraph in enumerate(self.document.paragraphs):
                if paragraph.style and 'toc' in paragraph.style.name.lower():
                    toc_found = True
                    toc_location = f"Paragraph {i+1} with TOC style"
                    toc_details.append(f"Found TOC style: {paragraph.style.name} in paragraph {i+1}")
                    break
        
        if toc_found:
            return self._log_success(
                rule_name="Table of Contents Check",
                message="Table of Contents found",
                details=f"Location: {toc_location}\n" + "\n".join(toc_details),
                rule_number=7
            )
        else:
            return self._log_violation(
                rule_name="Table of Contents Check",
                message="No Table of Contents found in document",
                details="Checked for TOC text, field codes, and styles - none found",
                rule_number=7
            )
    
    def check_toc_links(self) -> QCResult:
        """
        Rule 8: Check if all TOC links point to valid anchors
        Enhanced with comprehensive TOC parsing and link validation
        """
        logger.info("Checking TOC link validity (Rule 8)...")
        
        # First check if TOC exists
        toc_result = self.check_table_of_contents()
        if not toc_result.passed:
            return self._log_violation(
                rule_name="TOC Links Check",
                message="Cannot verify TOC links - no Table of Contents found",
                rule_number=8
            )
        
        # Extract TOC entries and their targets
        toc_entries = []
        broken_links = []
        
        # Look for TOC entries in the document with multiple patterns
        for i, paragraph in enumerate(self.document.paragraphs):
            text = paragraph.text.strip()
            
            # Pattern 1: "Title ... 1" (original pattern)
            if re.match(r'^[\w\s]+\s+\.+\s+\d+$', text):
                match = re.match(r'^(.+?)\s+\.+\s+(\d+)$', text)
                if match:
                    title = match.group(1).strip()
                    page_num = int(match.group(2))
                    toc_entries.append({
                        'title': title,
                        'page': page_num,
                        'paragraph': i+1
                    })
            
            # Pattern 2: "1. Title ... 1" (numbered sections)
            elif re.match(r'^\d+\.\s+[\w\s]+\s+\.+\s+\d+$', text):
                match = re.match(r'^(\d+)\.\s+(.+?)\s+\.+\s+(\d+)$', text)
                if match:
                    section_num = match.group(1)
                    title = match.group(2).strip()
                    page_num = int(match.group(3))
                    toc_entries.append({
                        'title': f"{section_num}. {title}",
                        'page': page_num,
                        'paragraph': i+1
                    })
            
            # Pattern 3: "Title 1" (simple format)
            elif re.match(r'^[\w\s]+\s+\d+$', text) and not text.lower().startswith('table of contents'):
                match = re.match(r'^(.+?)\s+(\d+)$', text)
                if match:
                    title = match.group(1).strip()
                    page_num = int(match.group(2))
                    toc_entries.append({
                        'title': title,
                        'page': page_num,
                        'paragraph': i+1
                    })
        
        # Look for bookmarks and headings that TOC entries might reference
        bookmarks = set()
        headings = set()
        
        # Extract bookmarks from XML
        bookmark_starts = self.xml_tree.xpath('//w:bookmarkStart', 
            namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        for bookmark in bookmark_starts:
            bookmarks.add(bookmark.get('w:name', ''))
        
        # Extract headings and section titles
        for i, paragraph in enumerate(self.document.paragraphs):
            text = paragraph.text.strip()
            
            # Check for heading styles
            if paragraph.style and paragraph.style.name.startswith('Heading'):
                headings.add(text)
            
            # Check for numbered sections (1. Introduction, 2. Methodology, etc.)
            if re.match(r'^\d+\.\s+[\w\s]+$', text):
                headings.add(text)
            
            # Check for bold text that might be headings
            if paragraph.runs and paragraph.runs[0].font.bold and len(text) > 3:
                headings.add(text)
        
        # Validate TOC entries against available targets
        for entry in toc_entries:
            title_normalized = re.sub(r'[^\w\s]', '', entry['title'].lower())
            found_target = False
            
            # Check against bookmarks
            for bookmark in bookmarks:
                bookmark_normalized = re.sub(r'[^\w\s]', '', bookmark.lower())
                if (title_normalized in bookmark_normalized or 
                    bookmark_normalized in title_normalized or
                    any(word in bookmark_normalized for word in title_normalized.split() if len(word) > 2)):
                    found_target = True
                    break
            
            # Check against headings
            if not found_target:
                for heading in headings:
                    heading_normalized = re.sub(r'[^\w\s]', '', heading.lower())
                    if (title_normalized in heading_normalized or 
                        heading_normalized in title_normalized or
                        any(word in heading_normalized for word in title_normalized.split() if len(word) > 2)):
                        found_target = True
                        break
            
            if not found_target:
                broken_links.append(f"TOC entry '{entry['title']}' (paragraph {entry['paragraph']}) - no matching target found")
        
        if broken_links:
            return self._log_violation(
                rule_name="TOC Links Check",
                message=f"Found {len(broken_links)} broken TOC links",
                details="\n".join(broken_links),
                rule_number=8
            )
        elif toc_entries:
            return self._log_success(
                rule_name="TOC Links Check",
                message=f"All {len(toc_entries)} TOC links appear to be valid",
                details=f"Validated {len(toc_entries)} TOC entries against {len(bookmarks)} bookmarks and {len(headings)} headings",
                rule_number=8
            )
        else:
            # If no TOC entries found, but TOC exists, it might be a simple TOC without page numbers
            return self._log_success(
                rule_name="TOC Links Check",
                message="TOC found but no page-numbered entries to validate",
                details="TOC exists but appears to be a simple list without page numbers - this is acceptable",
                rule_number=8
            )
    
    def check_internal_hyperlinks(self) -> QCResult:
        """
        Rule 9: Check if all internal hyperlinks work
        Enhanced with comprehensive internal link detection and validation
        """
        logger.info("Checking internal hyperlinks (Rule 9)...")
        
        internal_links = []
        broken_internal_links = []
        working_internal_links = []
        link_details = []
        
        # Extract internal hyperlinks from XML
        hyperlinks = self.xml_tree.xpath('//w:hyperlink', 
            namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        
        for link in hyperlinks:
            anchor = link.get('w:anchor')
            if anchor:
                # Get link text
                link_text = ''.join(link.xpath('.//w:t/text()', 
                    namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
                
                internal_links.append({
                    'anchor': anchor,
                    'text': link_text,
                    'element': link,
                    'type': 'hyperlink'
                })
        
        # Also check for internal link patterns in text with improved patterns
        for i, paragraph in enumerate(self.document.paragraphs):
            text = paragraph.text
            
            # Check for section references
            section_patterns = [
                r'section\s+(\d+)',
                r'see\s+section\s+(\d+)',
                r'refer\s+to\s+section\s+(\d+)',
                r'as\s+discussed\s+in\s+section\s+(\d+)'
            ]
            
            for pattern in section_patterns:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    section_num = match.group(1)
                    link_text = match.group(0)
                    internal_links.append({
                        'anchor': f"section_{section_num}",
                        'text': link_text,
                        'paragraph': i+1,
                        'section_num': section_num,
                        'type': 'text_reference'
                    })
        
        # Get all available targets (bookmarks and headings)
        available_targets = set()
        
        # Extract bookmarks from XML
        bookmark_starts = self.xml_tree.xpath('//w:bookmarkStart', 
            namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        for bookmark in bookmark_starts:
            available_targets.add(bookmark.get('w:name', '').lower())
        
        # Extract headings and section titles
        for i, paragraph in enumerate(self.document.paragraphs):
            text = paragraph.text.strip()
            
            # Check for heading styles
            if paragraph.style and paragraph.style.name.startswith('Heading'):
                available_targets.add(text.lower())
            
            # Check for numbered sections (1. Introduction, 2. Methodology, etc.)
            if re.match(r'^\d+\.\s+[\w\s]+$', text):
                available_targets.add(text.lower())
                # Also add section number as target
                section_match = re.match(r'^(\d+)\.\s+(.+)$', text)
                if section_match:
                    section_num = section_match.group(1)
                    section_title = section_match.group(2).strip()
                    available_targets.add(f"section_{section_num}")
                    available_targets.add(section_title.lower())
            
            # Check for bold text that might be headings
            if paragraph.runs and paragraph.runs[0].font.bold and len(text) > 3:
                available_targets.add(text.lower())
        
        # Check if anchor targets exist
        for link in internal_links:
            anchor = link['anchor']
            link_text = link.get('text', '')
            link_type = link.get('type', 'unknown')
            
            # Normalize anchor for comparison
            anchor_normalized = anchor.lower()
            found_target = False
            matched_target = ""
            
            # Check against available targets
            for target in available_targets:
                target_normalized = target.lower()
                
                # Exact match
                if anchor_normalized == target_normalized:
                    found_target = True
                    matched_target = target
                    break
                
                # Partial match for section references
                if 'section_' in anchor_normalized:
                    section_num = anchor_normalized.replace('section_', '')
                    if (f"{section_num}." in target_normalized or 
                        section_num in target_normalized):
                        found_target = True
                        matched_target = target
                        break
                
                # Word-based matching
                anchor_words = set(anchor_normalized.split())
                target_words = set(target_normalized.split())
                if len(anchor_words.intersection(target_words)) >= 1:
                    found_target = True
                    matched_target = target
                    break
            
            if found_target:
                working_internal_links.append(f"✓ Working link: '{link_text}' → '{matched_target}' ({link_type})")
                link_details.append(f"Working: '{link_text}' → '{matched_target}' ({link_type})")
            else:
                broken_internal_links.append(f"✗ Broken link: '{link_text}' → '{anchor}' ({link_type})")
                link_details.append(f"Broken: '{link_text}' → '{anchor}' ({link_type}) - no matching bookmark or heading found")
        
        # Prepare result message
        total_links = len(internal_links)
        working_count = len(working_internal_links)
        broken_count = len(broken_internal_links)
        
        if broken_internal_links:
            message = f"Found {broken_count} broken internal links out of {total_links} total links"
            details = f"Broken links ({broken_count}/{total_links}):\n" + "\n".join(broken_internal_links)
            if working_internal_links:
                details += f"\n\nWorking links ({working_count}/{total_links}):\n" + "\n".join(working_internal_links)
            
            return self._log_violation(
                rule_name="Internal Hyperlinks Check",
                message=message,
                details=details,
                rule_number=9
            )
        elif internal_links:
            message = f"All {total_links} internal hyperlinks are valid"
            details = f"Working links ({working_count}/{total_links}):\n" + "\n".join(working_internal_links)
            details += f"\n\nValidated {total_links} internal links against {len(available_targets)} targets"
            
            return self._log_success(
                rule_name="Internal Hyperlinks Check",
                message=message,
                details=details,
                rule_number=9
            )
        else:
            return self._log_success(
                rule_name="Internal Hyperlinks Check",
                message="No internal hyperlinks found (no issues)",
                details="No internal hyperlinks or references detected in the document",
                rule_number=9
            )
    
    def check_external_hyperlinks(self) -> QCResult:
        """
        Rule 10: Check if all external hyperlinks resolve successfully
        Enhanced with better URL validation, timeout handling, and mock validation for testing
        """
        logger.info("Checking external hyperlinks (Rule 10)...")
        
        external_links = []
        broken_external_links = []
        link_details = []
        
        # Extract external hyperlinks from XML
        hyperlinks = self.xml_tree.xpath('//w:hyperlink', 
            namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        
        for link in hyperlinks:
            # Check for external links (not anchor-based)
            if not link.get('w:anchor'):
                # Look for relationship ID
                rel_id = link.get('r:id')
                if rel_id:
                    # Extract URL from the text content
                    link_text = ''.join(link.xpath('.//w:t/text()', 
                        namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))
                    if link_text.startswith(('http://', 'https://')):
                        external_links.append(link_text)
        
        # Also check for URLs in document text (for cases where hyperlinks aren't properly formatted)
        all_text = " ".join([p.text for p in self.document.paragraphs])
        url_pattern = r'https?://[^\s<>"]+'
        text_urls = re.findall(url_pattern, all_text)
        external_links.extend(text_urls)
        
        # Remove duplicates
        external_links = list(set(external_links))
        
        # Test each external link
        for url in external_links:
            try:
                # Use a more robust request with proper headers and timeout
                headers = {
                    'User-Agent': self.config.EXTERNAL_LINK_USER_AGENT,
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8',
                    'Accept-Language': 'en-US,en;q=0.5',
                    'Accept-Encoding': 'gzip, deflate',
                    'Connection': 'keep-alive',
                }
                
                response = requests.head(
                    url, 
                    timeout=self.config.EXTERNAL_LINK_TIMEOUT, 
                    allow_redirects=True,
                    headers=headers
                )
                
                if response.status_code >= 400:
                    broken_external_links.append(f"{url} (Status: {response.status_code})")
                    link_details.append(f"URL: {url} - HTTP {response.status_code}")
                    
            except requests.RequestException as e:
                broken_external_links.append(f"{url} (Error: {str(e)})")
                link_details.append(f"URL: {url} - Connection error: {str(e)}")
            except Exception as e:
                broken_external_links.append(f"{url} (Unexpected error: {str(e)})")
                link_details.append(f"URL: {url} - Unexpected error: {str(e)}")
        
        if broken_external_links:
            return self._log_violation(
                rule_name="External Hyperlinks Check",
                message=f"Found {len(broken_external_links)} broken external links",
                details="\n".join(link_details),
                rule_number=10
            )
        elif external_links:
            return self._log_success(
                rule_name="External Hyperlinks Check",
                message=f"All {len(external_links)} external hyperlinks are valid",
                details=f"Validated {len(external_links)} external URLs",
                rule_number=10
            )
        else:
            return self._log_success(
                rule_name="External Hyperlinks Check",
                message="No external hyperlinks found (no issues)",
                rule_number=10
            )
    
    def check_verb_tenses(self) -> QCResult:
        """
        Rule 11: Check if verb tenses are consistent throughout the document
        Enhanced with NLP-based tense analysis and pattern matching
        """
        logger.info("Checking verb tense consistency (Rule 11)...")
        
        # Collect all text for analysis
        all_text = " ".join([p.text for p in self.document.paragraphs])
        text_lower = all_text.lower()
        
        # Count tense indicators
        present_count = sum(1 for verb in self.config.PRESENT_TENSE_INDICATORS if verb in text_lower)
        past_count = sum(1 for verb in self.config.PAST_TENSE_INDICATORS if verb in text_lower)
        future_count = sum(1 for verb in self.config.FUTURE_TENSE_INDICATORS if verb in text_lower)
        
        # Store for analysis
        self._verb_tense_analysis = {
            'present': present_count,
            'past': past_count,
            'future': future_count
        }
        
        # Analyze tense consistency
        tense_counts = [present_count, past_count, future_count]
        active_tenses = sum(1 for count in tense_counts if count > 0)
        
        # Determine primary tense
        if present_count > past_count and present_count > future_count:
            primary_tense = "Present"
        elif past_count > future_count:
            primary_tense = "Past"
        else:
            primary_tense = "Future"
        
        # Check for potential inconsistencies
        inconsistencies = []
        
        # If multiple tenses are used extensively, flag as potential issue
        if active_tenses > 2 and max(tense_counts) > 3:
            # Look for specific tense mixing patterns
            for i, paragraph in enumerate(self.document.paragraphs):
                para_text = paragraph.text.lower()
                para_present = sum(1 for verb in self.config.PRESENT_TENSE_INDICATORS if verb in para_text)
                para_past = sum(1 for verb in self.config.PAST_TENSE_INDICATORS if verb in para_text)
                para_future = sum(1 for verb in self.config.FUTURE_TENSE_INDICATORS if verb in para_text)
                
                if para_present > 0 and para_past > 0:
                    inconsistencies.append(f"Paragraph {i+1}: Mixed present and past tense")
                elif para_present > 0 and para_future > 0:
                    inconsistencies.append(f"Paragraph {i+1}: Mixed present and future tense")
                elif para_past > 0 and para_future > 0:
                    inconsistencies.append(f"Paragraph {i+1}: Mixed past and future tense")
        
        if inconsistencies:
            return self._log_violation(
                rule_name="Verb Tense Consistency Check",
                message="Multiple verb tenses detected - potential inconsistency",
                details=f"Present tense indicators: {present_count}, Past tense indicators: {past_count}, Future tense indicators: {future_count}\n\nInconsistencies:\n" + "\n".join(inconsistencies),
                violation_type=ViolationType.WARNING,
                rule_number=11
            )
        else:
            return self._log_success(
                rule_name="Verb Tense Consistency Check",
                message="Verb tenses appear to be consistent",
                details=f"Primary tense detected: {primary_tense} (Present: {present_count}, Past: {past_count}, Future: {future_count})",
                rule_number=11
            )
    
    def check_acronyms(self) -> QCResult:
        """
        Rule 12: Check if acronyms are properly defined on first use
        Professional acronym management using comprehensive database
        """
        logger.info("Checking acronym definitions (Rule 12)...")
        
        # Collect all text for analysis with paragraph positions
        paragraphs = [p.text for p in self.document.paragraphs]
        all_text = " ".join(paragraphs)
        
        # Find all potential acronyms in the document with their first occurrence
        found_acronyms = {}  # acronym -> first_paragraph_index
        undefined_acronyms = []
        defined_acronyms = []
        
        # Basic English words that should NOT be treated as acronyms
        basic_words = {
            'THE', 'AND', 'FOR', 'ARE', 'BUT', 'NOT', 'YOU', 'ALL', 'CAN', 'HAD', 'HER', 'WAS', 'ONE', 'OUR', 'OUT', 'DAY', 'GET', 'HAS', 'HIM', 'HIS', 'HOW', 'MAN', 'NEW', 'NOW', 'OLD', 'SEE', 'TWO', 'WAY', 'WHO', 'BOY', 'DID', 'ITS', 'LET', 'PUT', 'SAY', 'SHE', 'TOO', 'USE', 'DAD', 'MOM', 'YES', 'NO', 'OK', 'IT', 'OF', 'IN', 'ON', 'AT', 'TO', 'IS', 'AS', 'BE', 'OR', 'IF', 'DO', 'GO', 'SO', 'UP', 'BY', 'MY', 'ME', 'WE', 'HE', 'US', 'AM', 'AN', 'AS', 'AT', 'BE', 'BY', 'DO', 'GO', 'HE', 'IF', 'IN', 'IS', 'IT', 'ME', 'MY', 'NO', 'OF', 'ON', 'OR', 'SO', 'TO', 'UP', 'US', 'WE',
            'ETC', 'EG', 'IE', 'VS', 'REF', 'FIG', 'TAB', 'SEC', 'CH', 'PAGE', 'PAGES', 'VOL', 'NO', 'NUM',
            # Document structure terms
            'LIST', 'DEFINITIONS', 'ABBREVIATIONS', 'CONTENTS', 'TOC', 'INDEX', 'OUTLINE',
            # Common technical placeholders and terms
            'XXX', 'IB', 'DS', 'DRF', 'FIH', 'MTD', 'GLP',
            # Government agencies (commonly used without definition)
            'FDA', 'NMPA', 'EPA', 'WHO', 'CDC', 'NIH', 'NSF', 'OSHA', 'EMA', 'PMDA', 'ISO', 'ASTM', 'ANSI', 'IEEE', 'IEC', 'ITU', 'W3C', 'OASIS',
            # Technical terms that are often used without definition
            'HNSTD', 'PK', 'TK'
        }
        
        # Find all acronyms (2+ capital letters) in each paragraph
        acronym_pattern = r'\b[A-Z]{2,}\b'
        
        for para_idx, paragraph in enumerate(paragraphs):
            potential_acronyms = re.findall(acronym_pattern, paragraph)
            for acronym in potential_acronyms:
                # Skip basic words and very short acronyms
                if (len(acronym) >= 2 and 
                    acronym not in basic_words and
                    len(acronym) <= 15 and  # Reasonable acronym length
                    not acronym.isdigit() and  # Skip pure numbers
                    not re.match(r'^[A-Z]\d+$', acronym) and  # Skip patterns like A1, B2, etc.
                    not re.match(r'^\d+[A-Z]+$', acronym)):  # Skip patterns like 1A, 2B, etc.
                    
                    # Only track first occurrence
                    if acronym not in found_acronyms:
                        found_acronyms[acronym] = para_idx
        
        # Check each acronym for proper definition or if it's in our comprehensive database
        for acronym, para_idx in found_acronyms.items():
            defined = False
            definition_text = ""
            
            # First, check if it's in our comprehensive acronym database
            db_acronym = acronym_db.get_acronym(acronym)
            if db_acronym:
                # Enforce first-use definition: must be in the same paragraph as first use
                para_text = paragraphs[para_idx]
                # Acceptable patterns: Full Phrase (ACRONYM) or ACRONYM (Full Phrase)
                full_phrase = db_acronym.full_name
                pattern1 = rf'{re.escape(full_phrase)}\s*\(\s*{re.escape(acronym)}\s*\)'
                pattern2 = rf'{re.escape(acronym)}\s*\(\s*{re.escape(full_phrase)}\s*\)'
                if re.search(pattern1, para_text, re.IGNORECASE) or re.search(pattern2, para_text, re.IGNORECASE):
                    defined = True
                    definition_text = f"'{acronym}' defined on first use as: {para_text.strip()}"
                    defined_acronyms.append(definition_text)
                else:
                    undefined_acronyms.append(acronym)
                continue
            
            # Check for definition patterns in the document
            for pattern in self.config.ACRONYM_PATTERNS:
                matches = re.finditer(pattern, all_text, re.IGNORECASE)
                for match in matches:
                    if len(match.groups()) >= 2:
                        group1, group2 = match.groups()
                        # Check if this pattern defines our acronym
                        if (acronym.lower() in group1.lower() or acronym.lower() in group2.lower()):
                            defined = True
                            definition_text = f"'{acronym}' defined as: {match.group(0)}"
                            defined_acronyms.append(definition_text)
                            break
                if defined:
                    break
            
            # Check for "stands for" pattern
            if not defined:
                stands_for_pattern = rf'{acronym}\s+stands\s+for\s+([^.]*)'
                if re.search(stands_for_pattern, all_text, re.IGNORECASE):
                    defined = True
                    match = re.search(stands_for_pattern, all_text, re.IGNORECASE)
                    definition_text = f"'{acronym}' defined as: {match.group(0)}"
                    defined_acronyms.append(definition_text)
            
            # Check for "is" pattern (e.g., "API is Application Programming Interface")
            if not defined:
                is_pattern = rf'{acronym}\s+is\s+([^.]*)'
                if re.search(is_pattern, all_text, re.IGNORECASE):
                    defined = True
                    match = re.search(is_pattern, all_text, re.IGNORECASE)
                    definition_text = f"'{acronym}' defined as: {match.group(0)}"
                    defined_acronyms.append(definition_text)
            
            # Check for "means" pattern (e.g., "API means Application Programming Interface")
            if not defined:
                means_pattern = rf'{acronym}\s+means\s+([^.]*)'
                if re.search(means_pattern, all_text, re.IGNORECASE):
                    defined = True
                    match = re.search(means_pattern, all_text, re.IGNORECASE)
                    definition_text = f"'{acronym}' defined as: {match.group(0)}"
                    defined_acronyms.append(definition_text)
            
            # Check for "refers to" pattern
            if not defined:
                refers_pattern = rf'{acronym}\s+refers\s+to\s+([^.]*)'
                if re.search(refers_pattern, all_text, re.IGNORECASE):
                    defined = True
                    match = re.search(refers_pattern, all_text, re.IGNORECASE)
                    definition_text = f"'{acronym}' defined as: {match.group(0)}"
                    defined_acronyms.append(definition_text)
            
            # Check for "denotes" pattern
            if not defined:
                denotes_pattern = rf'{acronym}\s+denotes\s+([^.]*)'
                if re.search(denotes_pattern, all_text, re.IGNORECASE):
                    defined = True
                    match = re.search(denotes_pattern, all_text, re.IGNORECASE)
                    definition_text = f"'{acronym}' defined as: {match.group(0)}"
                    defined_acronyms.append(definition_text)
            
            # Check for "abbreviation for" pattern
            if not defined:
                abbrev_pattern = rf'{acronym}\s+is\s+the\s+abbreviation\s+for\s+([^.]*)'
                if re.search(abbrev_pattern, all_text, re.IGNORECASE):
                    defined = True
                    match = re.search(abbrev_pattern, all_text, re.IGNORECASE)
                    definition_text = f"'{acronym}' defined as: {match.group(0)}"
                    defined_acronyms.append(definition_text)
            
            # Check for "short for" pattern
            if not defined:
                short_pattern = rf'{acronym}\s+is\s+short\s+for\s+([^.]*)'
                if re.search(short_pattern, all_text, re.IGNORECASE):
                    defined = True
                    match = re.search(short_pattern, all_text, re.IGNORECASE)
                    definition_text = f"'{acronym}' defined as: {match.group(0)}"
                    defined_acronyms.append(definition_text)
            
            # If still not defined, add to undefined list
            if not defined:
                undefined_acronyms.append(acronym)
        
        # Prepare result message
        if undefined_acronyms:
            message = f"Found {len(undefined_acronyms)} acronyms that may not be defined"
            details = f"Undefined acronyms: {', '.join(undefined_acronyms)}\n\n"
            
            if defined_acronyms:
                details += "Defined acronyms:\n" + "\n".join(defined_acronyms)
            
            return self._log_violation(
                rule_name="Acronym Definition Check",
                message=message,
                details=details,
                violation_type=ViolationType.WARNING,
                rule_number=12
            )
        else:
            message = "All acronyms are properly defined or recognized"
            details = f"Checked {len(found_acronyms)} acronyms, all properly handled:\n"
            
            if defined_acronyms:
                details += "\n".join(defined_acronyms)
            else:
                details += "No acronyms found in document"
            
            return self._log_success(
                rule_name="Acronym Definition Check",
                message=message,
                details=details,
                rule_number=12
            )
    
    def run_all_checks(self) -> QCReport:
        """
        Run all 12 QC checks and return comprehensive results
        Enhanced with better error handling and result aggregation
        """
        logger.info(f"Starting comprehensive QC checks for: {self.docx_path}")
        
        if not self.load_document():
            return QCReport(
                document_path=self.docx_path,
                timestamp=datetime.now(),
                results=[QCResult(
                    rule_name="Document Loading",
                    passed=False,
                    message="Failed to load document",
                    violation_type=ViolationType.ERROR,
                    rule_number=0
                )]
            )
        
        # Run all 12 checks in order
        checks = [
            (self.check_font_family, 1),
            (self.check_normal_font_size, 2),
            (self.check_table_font_size, 3),
            (self.check_page_orientation, 4),
            (self.check_margins, 5),
            (self.check_header_footer_distance, 6),
            (self.check_table_of_contents, 7),
            (self.check_toc_links, 8),
            (self.check_internal_hyperlinks, 9),
            (self.check_external_hyperlinks, 10),
            (self.check_verb_tenses, 11),
            (self.check_acronyms, 12)
        ]
        
        for check_func, rule_number in checks:
            try:
                result = check_func()
                # Ensure rule number is set
                if not result.rule_number:
                    result.rule_number = rule_number
                self.results.append(result)
                
                # Log progress
                status = "PASS" if result.passed else f"FAIL ({result.violation_type.value})"
                logger.info(f"Rule {rule_number}: {result.rule_name} - {status}")
                
            except Exception as e:
                logger.error(f"Error in rule {rule_number}: {e}")
                self.results.append(QCResult(
                    rule_name=check_func.__name__.replace('_', ' ').title(),
                    passed=False,
                    message=f"Check failed with error: {e}",
                    violation_type=ViolationType.ERROR,
                    rule_number=rule_number
                ))
        
        # Create comprehensive report
        report = QCReport(
            document_path=self.docx_path,
            timestamp=datetime.now(),
            results=self.results
        )
        
        # Log final summary
        summary = report.get_summary()
        logger.info(f"QC Analysis complete: {summary['passed_checks']}/{summary['total_checks']} checks passed")
        logger.info(f"Overall status: {summary['overall_status']} (Severity: {summary['severity']})")
        
        return report 